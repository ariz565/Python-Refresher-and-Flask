"""
FastAPI + Vector Search Integration - Complete Production Guide
Build scalable vector search APIs with FastAPI
Covers embeddings, similarity search, and real-time applications
"""

print("=" * 80)
print("FASTAPI + VECTOR SEARCH - PRODUCTION READY APIs")
print("=" * 80)

# ============================================================================
# SECTION 1: FASTAPI FUNDAMENTALS FOR VECTOR SEARCH
# ============================================================================

print("\n" + "=" * 60)
print("SECTION 1: FASTAPI FUNDAMENTALS FOR VECTOR SEARCH")
print("=" * 60)

# Question 1: Why FastAPI for Vector Search?
print("\n1. Why use FastAPI for vector search applications?")
print("-" * 55)
print("""
FASTAPI ADVANTAGES FOR VECTOR SEARCH:

🚀 PERFORMANCE:
• Built on Starlette and Pydantic - extremely fast
• Async/await support for concurrent requests
• Automatic request/response validation
• Perfect for high-throughput similarity search

📊 API DESIGN:
• Automatic OpenAPI/Swagger documentation
• Type hints for better development experience
• Request/response models with validation
• Easy testing and debugging

🔧 VECTOR SEARCH SPECIFIC BENEFITS:
• Handles large request payloads (embeddings)
• Efficient JSON serialization for vectors
• Built-in error handling for malformed data
• Easy integration with ML libraries

⚡ SCALABILITY:
• Async processing for multiple searches
• Easy to deploy with Docker/Kubernetes
• Built-in dependency injection
• Middleware support for authentication/logging

COMMON VECTOR SEARCH API PATTERNS:
1. Text → Embedding → Search
2. Image → Embedding → Search  
3. Batch processing for multiple queries
4. Real-time indexing of new documents
5. Hybrid search (vector + traditional search)

FASTAPI vs ALTERNATIVES:
Framework  | Speed | Async | Docs | ML Integration
-----------|-------|-------|------|---------------
FastAPI    | ⭐⭐⭐⭐⭐ | ✅    | ⭐⭐⭐⭐⭐ | ⭐⭐⭐⭐⭐
Flask      | ⭐⭐⭐   | ❌    | ⭐⭐   | ⭐⭐⭐
Django     | ⭐⭐    | ⭐⭐⭐  | ⭐⭐⭐  | ⭐⭐
Tornado    | ⭐⭐⭐⭐  | ✅    | ⭐    | ⭐⭐

INTERVIEW QUESTIONS:
Q: "Why choose FastAPI over Flask for ML APIs?"
A: "FastAPI provides automatic validation, async support, 
   and better performance - crucial for vector search APIs
   that handle high-dimensional embeddings."

Q: "How does FastAPI handle large vector payloads?"
A: "FastAPI efficiently serializes JSON, supports streaming,
   and has built-in request size limits to prevent memory issues."
""")

# Question 2: Basic FastAPI Vector Search Setup
print("\n2. How to set up a basic FastAPI vector search API?")
print("-" * 54)
print("""
BASIC FASTAPI VECTOR SEARCH SETUP:

INSTALLATION:
pip install fastapi uvicorn numpy scikit-learn
pip install python-multipart  # For file uploads
pip install sentence-transformers  # For embeddings

BASIC PROJECT STRUCTURE:
vector_search_api/
├── main.py              # FastAPI app
├── models.py            # Pydantic models
├── vector_store.py      # Vector database
├── embeddings.py        # Embedding generation
├── requirements.txt     # Dependencies
└── config.py           # Configuration

CORE COMPONENTS:
1. Pydantic Models: Define request/response schemas
2. Vector Store: In-memory or persistent storage
3. Embedding Service: Convert text/images to vectors
4. Search Engine: Similarity search implementation
5. API Endpoints: RESTful endpoints for operations

BASIC API ENDPOINTS:
• POST /embed - Convert text to embeddings
• POST /search - Search similar vectors
• POST /index - Add new vectors to index
• GET /stats - Get index statistics
• DELETE /clear - Clear the index

EXAMPLE IMPLEMENTATION:
""")

def demonstrate_basic_fastapi_setup():
    """Demonstrate basic FastAPI vector search API setup"""
    print("Basic FastAPI Vector Search API Demo:")
    
    # Show the complete implementation
    print("\n1. main.py - FastAPI Application:")
    print("-" * 40)
    print("""
from fastapi import FastAPI, HTTPException
from typing import List, Optional
import uvicorn
from models import *
from vector_store import VectorStore
from embeddings import EmbeddingService

app = FastAPI(
    title="Vector Search API",
    description="High-performance similarity search API",
    version="1.0.0"
)

# Initialize services
vector_store = VectorStore()
embedding_service = EmbeddingService()

@app.post("/embed", response_model=EmbeddingResponse)
async def create_embedding(request: EmbeddingRequest):
    try:
        embedding = await embedding_service.embed(request.text)
        return EmbeddingResponse(
            text=request.text,
            embedding=embedding,
            model=embedding_service.model_name
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/search", response_model=SearchResponse)
async def search_similar(request: SearchRequest):
    try:
        # Get query embedding
        query_embedding = await embedding_service.embed(request.query)
        
        # Search similar vectors
        results = vector_store.search(
            query_embedding, 
            k=request.k,
            threshold=request.threshold
        )
        
        return SearchResponse(
            query=request.query,
            results=results,
            total_found=len(results)
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/index", response_model=IndexResponse)
async def add_to_index(request: IndexRequest):
    try:
        # Generate embedding
        embedding = await embedding_service.embed(request.text)
        
        # Add to vector store
        vector_id = vector_store.add(
            vector=embedding,
            metadata={
                "text": request.text,
                "category": request.category,
                "timestamp": request.timestamp or datetime.now()
            }
        )
        
        return IndexResponse(
            id=vector_id,
            text=request.text,
            indexed=True
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
    """)
    
    print("\n2. models.py - Pydantic Models:")
    print("-" * 35)
    print("""
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
from datetime import datetime

class EmbeddingRequest(BaseModel):
    text: str = Field(..., description="Text to embed")
    
class EmbeddingResponse(BaseModel):
    text: str
    embedding: List[float]
    model: str

class SearchRequest(BaseModel):
    query: str = Field(..., description="Search query")
    k: int = Field(default=5, ge=1, le=100, description="Number of results")
    threshold: Optional[float] = Field(default=None, description="Similarity threshold")

class SearchResult(BaseModel):
    id: str
    text: str
    score: float
    metadata: Dict[str, Any] = {}

class SearchResponse(BaseModel):
    query: str
    results: List[SearchResult]
    total_found: int

class IndexRequest(BaseModel):
    text: str = Field(..., description="Text to index")
    category: Optional[str] = None
    timestamp: Optional[datetime] = None

class IndexResponse(BaseModel):
    id: str
    text: str
    indexed: bool
    """)
    
    print("\n3. vector_store.py - Vector Storage:")
    print("-" * 38)
    print("""
import numpy as np
from typing import List, Dict, Any, Optional
import uuid
from sklearn.metrics.pairwise import cosine_similarity

class VectorStore:
    def __init__(self):
        self.vectors = []      # List of vectors
        self.metadata = []     # List of metadata
        self.ids = []          # List of IDs
    
    def add(self, vector: List[float], metadata: Dict[str, Any]) -> str:
        vector_id = str(uuid.uuid4())
        self.vectors.append(np.array(vector))
        self.metadata.append(metadata)
        self.ids.append(vector_id)
        return vector_id
    
    def search(self, query_vector: List[float], k: int = 5, 
               threshold: Optional[float] = None) -> List[Dict]:
        if not self.vectors:
            return []
        
        query_array = np.array(query_vector).reshape(1, -1)
        vectors_matrix = np.vstack(self.vectors)
        
        # Calculate cosine similarities
        similarities = cosine_similarity(query_array, vectors_matrix)[0]
        
        # Get top k indices
        top_indices = np.argsort(similarities)[::-1][:k]
        
        results = []
        for idx in top_indices:
            similarity = similarities[idx]
            
            # Apply threshold if specified
            if threshold and similarity < threshold:
                continue
                
            results.append({
                "id": self.ids[idx],
                "text": self.metadata[idx].get("text", ""),
                "score": float(similarity),
                "metadata": self.metadata[idx]
            })
        
        return results
    
    def get_stats(self) -> Dict[str, Any]:
        return {
            "total_vectors": len(self.vectors),
            "dimension": len(self.vectors[0]) if self.vectors else 0,
            "memory_usage_mb": sum(v.nbytes for v in self.vectors) / 1024 / 1024
        }
    """)
    
    print("\n4. embeddings.py - Embedding Service:")
    print("-" * 40)
    print("""
from sentence_transformers import SentenceTransformer
from typing import List
import asyncio

class EmbeddingService:
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.model_name = model_name
        self.model = SentenceTransformer(model_name)
    
    async def embed(self, text: str) -> List[float]:
        # Run embedding in thread to not block async loop
        loop = asyncio.get_event_loop()
        embedding = await loop.run_in_executor(
            None, 
            self._embed_sync, 
            text
        )
        return embedding.tolist()
    
    def _embed_sync(self, text: str):
        return self.model.encode(text)
    
    async def embed_batch(self, texts: List[str]) -> List[List[float]]:
        loop = asyncio.get_event_loop()
        embeddings = await loop.run_in_executor(
            None,
            self._embed_batch_sync,
            texts
        )
        return embeddings.tolist()
    
    def _embed_batch_sync(self, texts: List[str]):
        return self.model.encode(texts)
    """)

demonstrate_basic_fastapi_setup()

# ============================================================================
# SECTION 2: ADVANCED FASTAPI VECTOR SEARCH PATTERNS
# ============================================================================

print("\n" + "=" * 60)
print("SECTION 2: ADVANCED FASTAPI VECTOR SEARCH PATTERNS")
print("=" * 60)

# Question 3: Production-Ready Features
print("\n3. What production features should a vector search API include?")
print("-" * 69)
print("""
PRODUCTION-READY VECTOR SEARCH API FEATURES:

🔒 AUTHENTICATION & AUTHORIZATION:
• API key authentication
• Rate limiting per user/API key
• Role-based access control
• Request logging and audit trails

📊 MONITORING & OBSERVABILITY:
• Health checks and readiness probes
• Metrics collection (latency, throughput)
• Error tracking and alerting
• Performance monitoring

💾 DATA PERSISTENCE:
• Database integration (PostgreSQL + pgvector)
• Backup and recovery procedures
• Data versioning and migration
• Consistency guarantees

⚡ PERFORMANCE OPTIMIZATION:
• Caching frequently accessed vectors
• Connection pooling for databases
• Batch processing for multiple queries
• Async processing for heavy operations

🛡️ ERROR HANDLING & VALIDATION:
• Input validation and sanitization
• Graceful error responses
• Circuit breaker patterns
• Timeout handling

📈 SCALABILITY:
• Horizontal scaling with load balancers
• Database sharding strategies
• Microservices architecture
• Auto-scaling based on load

🔧 CONFIGURATION MANAGEMENT:
• Environment-based configuration
• Feature flags for A/B testing
• Dynamic configuration updates
• Secrets management

PRODUCTION CHECKLIST:
✅ Docker containerization
✅ CI/CD pipeline setup
✅ Environment separation (dev/staging/prod)
✅ Load testing and benchmarking
✅ Security scanning and compliance
✅ Documentation and API versioning
✅ Backup and disaster recovery
✅ Monitoring and alerting setup
""")

def demonstrate_production_api():
    """Demonstrate production-ready FastAPI vector search API"""
    print("Production-Ready Vector Search API Demo:")
    
    print("\n1. main.py - Production FastAPI App:")
    print("-" * 40)
    print("""
from fastapi import FastAPI, HTTPException, Depends, Security, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.trustedhost import TrustedHostMiddleware
import time
import logging
from contextlib import asynccontextmanager
import uvicorn

from config import settings
from vector_store import VectorStoreManager
from embeddings import EmbeddingService
from auth import verify_api_key
from monitoring import metrics_middleware
from models import *

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize services
vector_store_manager = VectorStoreManager()
embedding_service = EmbeddingService()
security = HTTPBearer()

# Startup/shutdown handlers
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    logger.info("Starting vector search API...")
    await vector_store_manager.initialize()
    await embedding_service.initialize()
    yield
    # Shutdown
    logger.info("Shutting down vector search API...")
    await vector_store_manager.close()

app = FastAPI(
    title="Production Vector Search API",
    description="Scalable, production-ready similarity search",
    version="2.0.0",
    lifespan=lifespan,
    docs_url="/docs" if settings.ENVIRONMENT != "production" else None
)

# Middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)

app.add_middleware(
    TrustedHostMiddleware,
    allowed_hosts=settings.ALLOWED_HOSTS
)

app.add_middleware(metrics_middleware)

# Dependency for authentication
async def get_current_user(
    credentials: HTTPAuthorizationCredentials = Security(security)
):
    user = await verify_api_key(credentials.credentials)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid API key",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return user

# Health check endpoints
@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": time.time(),
        "version": "2.0.0"
    }

@app.get("/ready")
async def readiness_check():
    # Check if services are ready
    vector_ready = await vector_store_manager.is_ready()
    embedding_ready = await embedding_service.is_ready()
    
    if not (vector_ready and embedding_ready):
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Service not ready"
        )
    
    return {"status": "ready"}

# Main API endpoints
@app.post("/api/v1/search", response_model=SearchResponse)
async def search_vectors(
    request: SearchRequest,
    user = Depends(get_current_user)
):
    start_time = time.time()
    
    try:
        # Rate limiting check
        if not await check_rate_limit(user.id, "search"):
            raise HTTPException(
                status_code=status.HTTP_429_TOO_MANY_REQUESTS,
                detail="Rate limit exceeded"
            )
        
        # Validate input
        if len(request.query.strip()) == 0:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Query cannot be empty"
            )
        
        # Get embedding
        query_embedding = await embedding_service.embed(
            request.query,
            user_id=user.id
        )
        
        # Search vectors
        results = await vector_store_manager.search(
            query_embedding=query_embedding,
            k=request.k,
            threshold=request.threshold,
            filters=request.filters,
            user_id=user.id
        )
        
        # Log request
        logger.info(f"Search request by user {user.id}: "
                   f"query='{request.query}', k={request.k}, "
                   f"results={len(results)}, "
                   f"time={time.time()-start_time:.3f}s")
        
        return SearchResponse(
            query=request.query,
            results=results,
            total_found=len(results),
            processing_time=time.time() - start_time
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Search error for user {user.id}: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Internal server error"
        )

@app.post("/api/v1/index", response_model=IndexResponse)
async def index_document(
    request: IndexRequest,
    user = Depends(get_current_user)
):
    try:
        # Validate permissions
        if not user.can_index:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Insufficient permissions"
            )
        
        # Generate embedding
        embedding = await embedding_service.embed(
            request.text,
            user_id=user.id
        )
        
        # Add to vector store
        doc_id = await vector_store_manager.add_document(
            text=request.text,
            embedding=embedding,
            metadata=request.metadata,
            user_id=user.id
        )
        
        return IndexResponse(
            id=doc_id,
            text=request.text,
            indexed=True
        )
        
    except Exception as e:
        logger.error(f"Indexing error for user {user.id}: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to index document"
        )

if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host=settings.HOST,
        port=settings.PORT,
        workers=settings.WORKERS,
        log_level=settings.LOG_LEVEL
    )
    """)
    
    print("\n2. config.py - Configuration Management:")
    print("-" * 43)
    print("""
from pydantic_settings import BaseSettings
from typing import List

class Settings(BaseSettings):
    # API Configuration
    HOST: str = "0.0.0.0"
    PORT: int = 8000
    WORKERS: int = 4
    ENVIRONMENT: str = "development"
    LOG_LEVEL: str = "info"
    
    # Security
    API_KEY_HEADER: str = "X-API-Key"
    ALLOWED_ORIGINS: List[str] = ["*"]
    ALLOWED_HOSTS: List[str] = ["*"]
    
    # Vector Store
    VECTOR_STORE_TYPE: str = "postgresql"  # or "faiss", "memory"
    DATABASE_URL: str = "postgresql://user:pass@localhost/vectordb"
    VECTOR_DIMENSION: int = 384
    
    # Embeddings
    EMBEDDING_MODEL: str = "all-MiniLM-L6-v2"
    EMBEDDING_CACHE_SIZE: int = 1000
    
    # Performance
    MAX_REQUEST_SIZE: int = 10 * 1024 * 1024  # 10MB
    REQUEST_TIMEOUT: int = 30
    RATE_LIMIT_REQUESTS: int = 100
    RATE_LIMIT_WINDOW: int = 60
    
    # Monitoring
    ENABLE_METRICS: bool = True
    METRICS_PORT: int = 9090
    
    class Config:
        env_file = ".env"

settings = Settings()
    """)
    
    print("\n3. auth.py - Authentication System:")
    print("-" * 38)
    print("""
import hashlib
import time
from typing import Optional
from dataclasses import dataclass
import redis

# Redis for rate limiting and caching
redis_client = redis.Redis(host='localhost', port=6379, db=0)

@dataclass
class User:
    id: str
    name: str
    can_index: bool = True
    can_search: bool = True
    rate_limit: int = 100  # requests per minute

# In production, store in database
API_KEYS = {
    "sk-test-12345": User(id="user1", name="Test User"),
    "sk-prod-67890": User(id="user2", name="Production User"),
}

async def verify_api_key(api_key: str) -> Optional[User]:
    # Hash API key for security
    key_hash = hashlib.sha256(api_key.encode()).hexdigest()
    
    # Check cache first
    cached_user = redis_client.get(f"user:{key_hash}")
    if cached_user:
        return User(**eval(cached_user))  # In production, use proper serialization
    
    # Verify API key
    user = API_KEYS.get(api_key)
    if user:
        # Cache for 5 minutes
        redis_client.setex(
            f"user:{key_hash}", 
            300, 
            str(user.__dict__)
        )
        return user
    
    return None

async def check_rate_limit(user_id: str, operation: str) -> bool:
    key = f"rate_limit:{user_id}:{operation}"
    current_minute = int(time.time() / 60)
    window_key = f"{key}:{current_minute}"
    
    current_count = redis_client.get(window_key)
    if current_count is None:
        current_count = 0
    else:
        current_count = int(current_count)
    
    user = next((u for u in API_KEYS.values() if u.id == user_id), None)
    if not user:
        return False
    
    if current_count >= user.rate_limit:
        return False
    
    # Increment counter
    redis_client.incr(window_key)
    redis_client.expire(window_key, 60)  # Expire after 1 minute
    
    return True
    """)

demonstrate_production_api()

# ============================================================================
# SECTION 3: REAL-WORLD APPLICATIONS
# ============================================================================

print("\n" + "=" * 60)
print("SECTION 3: REAL-WORLD APPLICATIONS")
print("=" * 60)

# Question 4: Complete Document Search System
print("\n4. How to build a complete document search system?")
print("-" * 55)
print("""
COMPLETE DOCUMENT SEARCH SYSTEM ARCHITECTURE:

COMPONENTS:
1. Document Ingestion Pipeline
2. Text Preprocessing and Chunking
3. Embedding Generation Service
4. Vector Database (PostgreSQL + pgvector)
5. Search API with FastAPI
6. Frontend Interface
7. Monitoring and Analytics

SYSTEM FLOW:
Documents → Preprocessing → Chunking → Embeddings → Vector DB → Search API

KEY FEATURES:
✅ Multi-format support (PDF, Word, HTML, etc.)
✅ Intelligent text chunking
✅ Metadata preservation
✅ Semantic and keyword search
✅ Real-time indexing
✅ Search result ranking
✅ User feedback and analytics

TECHNICAL STACK:
• FastAPI: REST API framework
• PostgreSQL + pgvector: Vector storage
• Sentence Transformers: Embedding generation
• PyPDF2/python-docx: Document parsing
• Redis: Caching and rate limiting
• Docker: Containerization
• Nginx: Load balancing
""")

def demonstrate_document_search_system():
    """Demonstrate complete document search system"""
    print("Complete Document Search System Implementation:")
    
    print("\n1. Document Processing Pipeline:")
    print("-" * 40)
    print("""
import asyncio
import hashlib
from typing import List, Dict, Any
from dataclasses import dataclass
import PyPDF2
import docx
from io import BytesIO

@dataclass
class Document:
    id: str
    title: str
    content: str
    metadata: Dict[str, Any]
    chunks: List[str] = None

class DocumentProcessor:
    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    async def process_file(self, file_content: bytes, 
                          filename: str, metadata: Dict = None) -> Document:
        # Extract text based on file type
        if filename.endswith('.pdf'):
            content = await self._extract_pdf_text(file_content)
        elif filename.endswith('.docx'):
            content = await self._extract_docx_text(file_content)
        elif filename.endswith('.txt'):
            content = file_content.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file type: {filename}")
        
        # Generate document ID
        doc_id = hashlib.md5(
            (filename + content[:100]).encode()
        ).hexdigest()
        
        # Create document
        doc = Document(
            id=doc_id,
            title=filename,
            content=content,
            metadata=metadata or {}
        )
        
        # Chunk the content
        doc.chunks = self._chunk_text(content)
        
        return doc
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        reader = PyPDF2.PdfReader(BytesIO(content))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\\n"
        return text
    
    async def _extract_docx_text(self, content: bytes) -> str:
        doc = docx.Document(BytesIO(content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\\n"
        return text
    
    def _chunk_text(self, text: str) -> List[str]:
        # Simple sentence-based chunking
        sentences = text.split('. ')
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < self.chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    """)
    
    print("\n2. Enhanced Vector Store with Metadata:")
    print("-" * 44)
    print("""
import asyncpg
import numpy as np
from typing import List, Dict, Any, Optional
import json

class PostgreSQLVectorStore:
    def __init__(self, connection_string: str):
        self.connection_string = connection_string
        self.pool = None
    
    async def initialize(self):
        self.pool = await asyncpg.create_pool(self.connection_string)
        
        # Create tables
        async with self.pool.acquire() as conn:
            await conn.execute('''
                CREATE EXTENSION IF NOT EXISTS vector;
                
                CREATE TABLE IF NOT EXISTS documents (
                    id VARCHAR PRIMARY KEY,
                    title VARCHAR NOT NULL,
                    content TEXT NOT NULL,
                    metadata JSONB,
                    created_at TIMESTAMP DEFAULT NOW()
                );
                
                CREATE TABLE IF NOT EXISTS document_chunks (
                    id SERIAL PRIMARY KEY,
                    document_id VARCHAR REFERENCES documents(id),
                    chunk_index INTEGER,
                    content TEXT NOT NULL,
                    embedding VECTOR(384),
                    metadata JSONB,
                    created_at TIMESTAMP DEFAULT NOW()
                );
                
                CREATE INDEX IF NOT EXISTS idx_chunks_embedding 
                ON document_chunks USING ivfflat (embedding vector_cosine_ops)
                WITH (lists = 100);
            ''')
    
    async def add_document(self, document: Document, 
                          embeddings: List[List[float]]) -> str:
        async with self.pool.acquire() as conn:
            async with conn.transaction():
                # Insert document
                await conn.execute('''
                    INSERT INTO documents (id, title, content, metadata)
                    VALUES ($1, $2, $3, $4)
                    ON CONFLICT (id) DO UPDATE SET
                    title = EXCLUDED.title,
                    content = EXCLUDED.content,
                    metadata = EXCLUDED.metadata
                ''', document.id, document.title, document.content,
                    json.dumps(document.metadata))
                
                # Insert chunks with embeddings
                for i, (chunk, embedding) in enumerate(zip(document.chunks, embeddings)):
                    await conn.execute('''
                        INSERT INTO document_chunks 
                        (document_id, chunk_index, content, embedding, metadata)
                        VALUES ($1, $2, $3, $4, $5)
                    ''', document.id, i, chunk, embedding,
                        json.dumps({"chunk_index": i}))
        
        return document.id
    
    async def search(self, query_embedding: List[float], k: int = 10,
                    threshold: float = 0.7, filters: Dict = None) -> List[Dict]:
        
        where_clause = ""
        params = [query_embedding, k]
        
        if filters:
            if 'document_ids' in filters:
                where_clause += " AND dc.document_id = ANY($3)"
                params.append(filters['document_ids'])
        
        query = f'''
            SELECT 
                dc.document_id,
                dc.content,
                dc.metadata,
                d.title,
                d.metadata as doc_metadata,
                1 - (dc.embedding <=> $1) as similarity
            FROM document_chunks dc
            JOIN documents d ON dc.document_id = d.id
            WHERE 1 - (dc.embedding <=> $1) > {threshold}
            {where_clause}
            ORDER BY dc.embedding <=> $1
            LIMIT $2
        '''
        
        async with self.pool.acquire() as conn:
            rows = await conn.fetch(query, *params)
        
        results = []
        for row in rows:
            results.append({
                'document_id': row['document_id'],
                'title': row['title'],
                'content': row['content'],
                'similarity': float(row['similarity']),
                'metadata': row['metadata'],
                'doc_metadata': row['doc_metadata']
            })
        
        return results
    
    async def get_document_stats(self) -> Dict[str, Any]:
        async with self.pool.acquire() as conn:
            doc_count = await conn.fetchval('SELECT COUNT(*) FROM documents')
            chunk_count = await conn.fetchval('SELECT COUNT(*) FROM document_chunks')
            
        return {
            'total_documents': doc_count,
            'total_chunks': chunk_count,
            'avg_chunks_per_doc': chunk_count / max(doc_count, 1)
        }
    """)
    
    print("\n3. Document Search API Endpoints:")
    print("-" * 38)
    print("""
from fastapi import FastAPI, UploadFile, File, Form
from typing import List, Optional

@app.post("/api/v1/documents/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    category: Optional[str] = Form(None),
    tags: Optional[str] = Form(None),
    user = Depends(get_current_user)
):
    try:
        # Validate file type
        if not file.filename.endswith(('.pdf', '.docx', '.txt')):
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type"
            )
        
        # Read file content
        content = await file.read()
        
        # Process document
        metadata = {
            'category': category,
            'tags': tags.split(',') if tags else [],
            'uploaded_by': user.id,
            'file_size': len(content)
        }
        
        document = await document_processor.process_file(
            content, file.filename, metadata
        )
        
        # Generate embeddings for chunks
        embeddings = await embedding_service.embed_batch(document.chunks)
        
        # Store in vector database
        doc_id = await vector_store.add_document(document, embeddings)
        
        return DocumentUploadResponse(
            document_id=doc_id,
            title=document.title,
            chunks_created=len(document.chunks),
            status="indexed"
        )
        
    except Exception as e:
        logger.error(f"Document upload error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Failed to process document"
        )

@app.post("/api/v1/documents/search", response_model=DocumentSearchResponse)
async def search_documents(
    request: DocumentSearchRequest,
    user = Depends(get_current_user)
):
    try:
        # Generate query embedding
        query_embedding = await embedding_service.embed(request.query)
        
        # Search with filters
        filters = {}
        if request.document_ids:
            filters['document_ids'] = request.document_ids
        
        results = await vector_store.search(
            query_embedding=query_embedding,
            k=request.k,
            threshold=request.threshold,
            filters=filters
        )
        
        # Group results by document
        document_results = {}
        for result in results:
            doc_id = result['document_id']
            if doc_id not in document_results:
                document_results[doc_id] = {
                    'document_id': doc_id,
                    'title': result['title'],
                    'chunks': [],
                    'max_similarity': 0,
                    'metadata': result['doc_metadata']
                }
            
            document_results[doc_id]['chunks'].append({
                'content': result['content'],
                'similarity': result['similarity']
            })
            
            document_results[doc_id]['max_similarity'] = max(
                document_results[doc_id]['max_similarity'],
                result['similarity']
            )
        
        # Sort by max similarity
        sorted_docs = sorted(
            document_results.values(),
            key=lambda x: x['max_similarity'],
            reverse=True
        )
        
        return DocumentSearchResponse(
            query=request.query,
            documents=sorted_docs,
            total_found=len(sorted_docs)
        )
        
    except Exception as e:
        logger.error(f"Document search error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Search failed"
        )

@app.get("/api/v1/documents/{document_id}")
async def get_document(
    document_id: str,
    user = Depends(get_current_user)
):
    # Retrieve full document details
    pass

@app.delete("/api/v1/documents/{document_id}")
async def delete_document(
    document_id: str,
    user = Depends(get_current_user)
):
    # Delete document and its chunks
    pass
    """)

demonstrate_document_search_system()

# ============================================================================
# SECTION 4: DEPLOYMENT AND MONITORING
# ============================================================================

print("\n" + "=" * 60)
print("SECTION 4: DEPLOYMENT AND MONITORING")
print("=" * 60)

# Question 5: Production Deployment
print("\n5. How to deploy FastAPI vector search to production?")
print("-" * 57)
print("""
PRODUCTION DEPLOYMENT GUIDE:

🐳 CONTAINERIZATION:
• Docker for consistent environments
• Multi-stage builds for optimization
• Health checks and graceful shutdown
• Resource limits and monitoring

☸️ ORCHESTRATION:
• Kubernetes for scalability
• Load balancing across instances
• Auto-scaling based on metrics
• Rolling deployments with zero downtime

🔧 INFRASTRUCTURE:
• Database clustering for high availability
• Redis cluster for caching
• Nginx/Traefik for load balancing
• SSL/TLS termination

📊 MONITORING:
• Prometheus for metrics collection
• Grafana for visualization
• Structured logging with ELK stack
• Alert manager for notifications

🔒 SECURITY:
• API key management
• Network policies
• Secret management (Vault)
• Regular security scans

DEPLOYMENT CHECKLIST:
✅ Environment configuration
✅ Database migrations
✅ SSL certificate setup
✅ Monitoring configuration
✅ Backup procedures
✅ Load testing completion
✅ Security audit passed
✅ Documentation updated
""")

def demonstrate_production_deployment():
    """Demonstrate production deployment setup"""
    print("Production Deployment Configuration:")
    
    print("\n1. Dockerfile - Multi-stage Build:")
    print("-" * 38)
    print("""
# Dockerfile
FROM python:3.11-slim as builder

WORKDIR /app

# Install system dependencies
RUN apt-get update && apt-get install -y \\
    gcc \\
    g++ \\
    && rm -rf /var/lib/apt/lists/*

# Copy requirements and install Python dependencies
COPY requirements.txt .
RUN pip install --no-cache-dir --user -r requirements.txt

# Production stage
FROM python:3.11-slim

WORKDIR /app

# Install runtime dependencies
RUN apt-get update && apt-get install -y \\
    && rm -rf /var/lib/apt/lists/*

# Copy Python packages from builder
COPY --from=builder /root/.local /root/.local

# Copy application code
COPY . .

# Create non-root user
RUN useradd --create-home --shell /bin/bash app \\
    && chown -R app:app /app
USER app

# Add local Python packages to PATH
ENV PATH=/root/.local/bin:$PATH

# Health check
HEALTHCHECK --interval=30s --timeout=30s --start-period=5s --retries=3 \\
    CMD curl -f http://localhost:8000/health || exit 1

EXPOSE 8000

CMD ["uvicorn", "main:app", "--host", "0.0.0.0", "--port", "8000"]
    """)
    
    print("\n2. docker-compose.yml - Complete Stack:")
    print("-" * 42)
    print("""
version: '3.8'

services:
  vector-api:
    build: .
    ports:
      - "8000:8000"
    environment:
      - DATABASE_URL=postgresql://postgres:password@postgres:5432/vectordb
      - REDIS_URL=redis://redis:6379/0
      - ENVIRONMENT=production
    depends_on:
      - postgres
      - redis
    deploy:
      replicas: 3
      resources:
        limits:
          memory: 1G
          cpus: '0.5'
    restart: unless-stopped

  postgres:
    image: pgvector/pgvector:pg15
    environment:
      - POSTGRES_DB=vectordb
      - POSTGRES_USER=postgres
      - POSTGRES_PASSWORD=password
    volumes:
      - postgres_data:/var/lib/postgresql/data
      - ./init.sql:/docker-entrypoint-initdb.d/init.sql
    ports:
      - "5432:5432"
    deploy:
      resources:
        limits:
          memory: 2G
          cpus: '1.0'

  redis:
    image: redis:7-alpine
    ports:
      - "6379:6379"
    volumes:
      - redis_data:/data
    command: redis-server --appendonly yes

  nginx:
    image: nginx:alpine
    ports:
      - "80:80"
      - "443:443"
    volumes:
      - ./nginx.conf:/etc/nginx/nginx.conf
      - ./ssl:/etc/nginx/ssl
    depends_on:
      - vector-api

  prometheus:
    image: prom/prometheus
    ports:
      - "9090:9090"
    volumes:
      - ./prometheus.yml:/etc/prometheus/prometheus.yml
      - prometheus_data:/prometheus

  grafana:
    image: grafana/grafana
    ports:
      - "3000:3000"
    environment:
      - GF_SECURITY_ADMIN_PASSWORD=admin
    volumes:
      - grafana_data:/var/lib/grafana

volumes:
  postgres_data:
  redis_data:
  prometheus_data:
  grafana_data:
    """)
    
    print("\n3. Kubernetes Deployment:")
    print("-" * 32)
    print("""
# k8s-deployment.yaml
apiVersion: apps/v1
kind: Deployment
metadata:
  name: vector-search-api
spec:
  replicas: 3
  selector:
    matchLabels:
      app: vector-search-api
  template:
    metadata:
      labels:
        app: vector-search-api
    spec:
      containers:
      - name: api
        image: your-registry/vector-search-api:latest
        ports:
        - containerPort: 8000
        env:
        - name: DATABASE_URL
          valueFrom:
            secretKeyRef:
              name: database-secret
              key: url
        resources:
          requests:
            memory: "512Mi"
            cpu: "250m"
          limits:
            memory: "1Gi"
            cpu: "500m"
        livenessProbe:
          httpGet:
            path: /health
            port: 8000
          initialDelaySeconds: 30
          periodSeconds: 10
        readinessProbe:
          httpGet:
            path: /ready
            port: 8000
          initialDelaySeconds: 5
          periodSeconds: 5

---
apiVersion: v1
kind: Service
metadata:
  name: vector-search-service
spec:
  selector:
    app: vector-search-api
  ports:
  - port: 80
    targetPort: 8000
  type: LoadBalancer

---
apiVersion: networking.k8s.io/v1
kind: Ingress
metadata:
  name: vector-search-ingress
  annotations:
    kubernetes.io/ingress.class: nginx
    cert-manager.io/cluster-issuer: letsencrypt-prod
spec:
  tls:
  - hosts:
    - api.yourdomain.com
    secretName: vector-search-tls
  rules:
  - host: api.yourdomain.com
    http:
      paths:
      - path: /
        pathType: Prefix
        backend:
          service:
            name: vector-search-service
            port:
              number: 80
    """)

demonstrate_production_deployment()

print("\n" + "=" * 80)
print("VECTOR SEARCH API INTERVIEWS - COMMON QUESTIONS")
print("=" * 80)

print("""
TOP INTERVIEW QUESTIONS & ANSWERS:

Q1: "How would you handle embedding model updates in production?"
A: "Use versioned embeddings, gradual rollout, A/B testing, and 
   maintain backward compatibility during transition periods."

Q2: "What's your strategy for scaling vector search to millions of vectors?"
A: "Implement database sharding, use approximate algorithms (LSH/HNSW),
   add caching layers, and consider specialized vector databases."

Q3: "How do you ensure search result quality?"
A: "Implement relevance feedback, A/B test different embedding models,
   use hybrid search (vector + keyword), and monitor search metrics."

Q4: "Describe your approach to handling cold starts."
A: "Preload popular embeddings, implement warming strategies,
   use fallback to traditional search, and optimize model loading."

Q5: "How would you implement real-time indexing?"
A: "Use async processing, implement batch insertions, 
   maintain separate hot/cold indices, and use CDC for updates."

KEY TAKEAWAYS:
✅ FastAPI + Vector Search = High-performance semantic search APIs
✅ Production requires monitoring, authentication, and scalability
✅ Choose appropriate storage (pgvector vs FAISS vs specialized DBs)
✅ Balance accuracy vs speed based on use case requirements
✅ Implement proper error handling and graceful degradation

NEXT STEPS:
🔗 Build a prototype with your own data
🔗 Experiment with different embedding models
🔗 Practice deployment with Docker/Kubernetes
🔗 Implement monitoring and observability
🔗 Study specialized vector databases (Pinecone, Weaviate, etc.)
""")

print("\n" + "=" * 80)
print("END OF FASTAPI + VECTOR SEARCH GUIDE")
print("You now have complete knowledge for building production vector search APIs!")
print("=" * 80)
